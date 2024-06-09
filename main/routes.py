from datetime import datetime
from flask import render_template, url_for, flash, redirect, request, abort
from main import app, db, bcrypt
from main.forms import RegistrationForm, LoginForm, UpdateAccountForm, PostForm,TestForm
from main.models import User, Post, Publication, PassedTests
from flask_login import login_user, current_user, logout_user, login_required
import os
import secrets
from PIL import Image
import json
from flask import send_file, make_response
import tempfile
import docx
from docx.shared import Pt


@app.route("/")
def lending():
    return render_template('landing.html', title='Сетьстрой обучение')

@app.route("/home")
def home():
    page = request.args.get('page', 1, type=int)
    posts = Post.query.order_by(Post.date_posted.desc()).paginate(page=page, per_page=6)
    return render_template('home.html', posts=posts)


@app.route("/about")
def about():
    return render_template('about.html', title='Подробнее')


@app.route("/register", methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    form = RegistrationForm()
    if form.validate_on_submit():
        hashed_password = bcrypt.generate_password_hash(form.password.data).decode('utf-8')
        user = User(username=form.username.data, email=form.email.data, password=hashed_password)
        db.session.add(user)
        db.session.commit()
        flash('Ваш аккаунт был создан! Теперь вы можете войти в систему', 'success')
        return redirect(url_for('login'))
    return render_template('register.html', title='Регистрация', form=form)


@app.route("/login", methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('home'))
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(email=form.email.data).first()
        if user and bcrypt.check_password_hash(user.password, form.password.data):
            login_user(user, remember=form.remember.data)
            next_page = request.args.get('next')
            return redirect(next_page) if next_page else redirect(url_for('home'))
        else:
            flash('Вход неуспешный. Пожалуйста, проверьте электронную почту и пароль', 'danger')
    return render_template('login.html', title='Войти', form=form)


@app.route("/logout")
def logout():
    logout_user()
    return redirect(url_for('home'))


def save_picture(form_picture):
    random_hex = secrets.token_hex(8)
    _, f_ext = os.path.splitext(form_picture.filename)
    picture_fn = random_hex + f_ext
    picture_path = os.path.join(app.root_path, 'static/profile_pics', picture_fn)

    output_size = (125, 125)
    i = Image.open(form_picture)
    i.thumbnail(output_size)
    i.save(picture_path)

    return picture_fn


@app.route("/account", methods=['GET', 'POST'])
@login_required
def account():
    form = UpdateAccountForm()
    if form.validate_on_submit():
        if form.picture.data:
            picture_file = save_picture(form.picture.data)
            current_user.image_file = picture_file
        current_user.username = form.username.data
        current_user.email = form.email.data
        db.session.commit()
        flash('Ваш аккаунт был обновлен!', 'success')
        return redirect(url_for('account'))
    elif request.method == 'GET':
        form.username.data = current_user.username
        form.email.data = current_user.email
    image_file = url_for('static', filename='profile_pics/' + current_user.image_file)
    return render_template('account.html', title='Аккаунт',
                           image_file=image_file, form=form)


@app.route("/post/new", methods=['GET', 'POST'])
@login_required
def new_post():
    form = PostForm()
    if form.validate_on_submit():
        post = Post(title=form.title.data, content=form.content.data, author=current_user)
        db.session.add(post)
        db.session.commit()
        flash('Ваш материал создан!', 'success')
        return redirect(url_for('home'))
    return render_template('create_post.html', title='Новая запись',
                           form=form, legend='Новая публикация')


@app.route("/post/<int:post_id>")
def post(post_id):
    post = Post.query.get_or_404(post_id)
    quest = db.session.query(Publication).filter(Publication.post_id == post_id).all()
    if quest:
        quest = quest[0]
    else:
        quest = None
    return render_template('post.html', quest=quest, title=post.title, post=post)


@app.route("/post/<int:post_id>/update", methods=['GET', 'POST'])
@login_required
def update_post(post_id):
    post = Post.query.get_or_404(post_id)
    if post.author != current_user:
        abort(403)
    form = PostForm()
    if form.validate_on_submit():
        post.title = form.title.data
        post.content = form.content.data
        db.session.commit()
        flash('Ваш материал изменен!', 'success')
        return redirect(url_for('post', post_id=post.id))
    elif request.method == 'GET':
        form.title.data = post.title
        form.content.data = post.content
    return render_template('create_post.html', title='Изменить статью',
                           form=form, legend='Изменить статью')


@app.route("/post/<int:post_id>/delete", methods=['POST'])
@login_required
def delete_post(post_id):
    post = Post.query.get_or_404(post_id)
    if post.author != current_user:
        abort(403)
    db.session.delete(post)
    db.session.commit()
    flash('Ваш материал удален!', 'success')
    return redirect(url_for('home'))


@app.route("/user/<string:username>")
def user_posts(username):
    page = request.args.get('page', 1, type=int)
    user = User.query.filter_by(username=username).first_or_404()
    posts = Post.query.filter_by(author=user)\
        .order_by(Post.date_posted.desc())\
        .paginate(page=page, per_page=6)
    return render_template('user_posts.html', posts=posts, user=user)


@app.route("/post/<int:post_id>/test", methods=['GET', 'POST'])
@login_required
def post_test(post_id):
    post = Post.query.get_or_404(post_id)
    is_not_author = post.author != current_user

    quest = db.session.query(Publication).filter(Publication.post_id == post_id).all()
    if quest:
        quest = quest[0]
        questions = json.loads(quest.questions)
    else:
        questions = 'none'
    return render_template('test.html', question=questions, quest=quest ,title='Тестирование', is_not_author=is_not_author)


@app.route("/post/<int:post_id>/test/change", methods=['GET', 'POST'])
@login_required
def change_test(post_id):
    post = Post.query.get_or_404(post_id)
    if post.author != current_user:
        abort(403)
    quest = db.session.query(Publication).filter(Publication.post_id == post_id).all()

    if quest:
        quest = quest[0]
        form = TestForm()
        if form.validate_on_submit():
            quest.title = form.title.data
            quest.questions = form.questions.data
            db.session.commit()
            flash('Ваш материал изменен!', 'success')
            return redirect(url_for('post_test', post_id=quest.post_id))
        elif request.method == 'GET':
            form.title.data = quest.title
            form.questions.data = quest.questions
    else:
        form = 'none'

    return render_template('change_test.html', form=form, quest=quest, title='Редактор теста', legend='Изменить тест')


@app.route("/post/<int:post_id>/test/new", methods=['GET', 'POST'])
@login_required
def new_test(post_id):
    form = TestForm()
    if form.validate_on_submit():
        publication = Publication(title=form.title.data, questions=form.questions.data, post_id=post_id)
        db.session.add(publication)
        db.session.commit()
        flash('Ваш тест создан!', 'success')
        return redirect(url_for('home'))
    return render_template('new_test.html', title='Новый тест',
                           form=form, legend='Новый тест')


@app.route("/post/<int:post_id>/test/result", methods=['GET', 'POST'])
@login_required
def result(post_id):
    a = request.form

    quest = db.session.query(Publication).filter(Publication.post_id == post_id).all()
    if quest:
        quest = quest[0]
        questions = json.loads(quest.questions)
    else:
        quest = 'none'
        questions = 'none'
    dt = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    return render_template('result.html', title='Новый тест',post_id=post_id, questions=questions,quest=quest, legend='Новый тест', data=a, dt=dt)


def save_word(passed_test):
    user_by_id = db.session.query(User).filter(User.id == passed_test.user_id).order_by(User.id.desc()).first()
    publication_by_id = db.session.query(Publication).filter(Publication.id == passed_test.test_id).order_by(Publication.id.desc()).first()

    doc = docx.Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(14)
    doc.add_paragraph('Название тестирования: ' + publication_by_id.title)
    doc.add_paragraph('Прошел тестирование: ' + user_by_id.username)
    doc.add_paragraph('Процент правильных ответов: ' + str(passed_test.score) + '%')
    doc.add_paragraph('Дата прохождения тестирования: ' + passed_test.passe_date.strftime("%d.%m.%Y %H:%M:%S"))

    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp:
        doc.save(temp.name)
        temp_path = temp.name
    
    return temp_path


@app.route("/post/<int:post_id>/test/result_handling/<int:score>/<int:user_id>/<int:test_id>/<string:passe_date>", methods=['GET', 'POST'])
@login_required
def result_handling(post_id, score, user_id, test_id, passe_date):
    try:
        a = datetime.strptime(passe_date, "%d.%m.%Y %H:%M:%S")
        result2 = PassedTests(score=score, passe_date=a, user_id=user_id, test_id=test_id)
        db.session.add(result2)
        db.session.commit()
        passed_test = db.session.query(PassedTests).filter(PassedTests.user_id == current_user.id
                                                           and PassedTests.test_id == test_id).order_by(PassedTests.id.desc()).first()
        temp_path = save_word(passed_test)
        flash('Тестирование сохранено!', 'success')
    except Exception as e:
        flash(f'Тестирование не сохранено! Ошибка: {str(e)}', 'danger')
        return redirect(url_for('home'))
    
    response = make_response(send_file(temp_path, as_attachment=True))
    response.headers['Content-Disposition'] = f'attachment; filename=test_result_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    return response

@app.errorhandler(403)
@app.errorhandler(404)
def pageNotFound(error):
    return render_template('page404.html', title='Страница не найдена')
